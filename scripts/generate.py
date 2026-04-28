"""457deep 경력기술서 자동 생성 스크립트.

흐름 (skill.md PART A):
1. outputs/{오늘날짜}/ 가 이미 있으면 즉시 종료 (덮어쓰기 안 함)
2. history.json 읽기 (없으면 초기화)
3. remaining_in_cycle 에서 직무 3개 무작위 선택
4. 각 직무에 (연차, 이직 상황) 조합 배정 (used_combos 회피)
5. 각 직무에 양식 2개 배정 (직무·연차 매핑 규칙 + 최근 5일 페어 회피)
6. 각 직무마다 Claude API 1회 호출 → 가상 인물 + 양식 2개 본문 생성
7. 6개 docx 저장
8. history.json 업데이트
워크플로우(generate.yml)가 outputs/, history.json 을 git 커밋·푸시.
"""

from __future__ import annotations

import json
import os
import random
import re
import sys
from datetime import datetime, timedelta, timezone
from pathlib import Path

import anthropic
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parent.parent
SKILL_PATH = ROOT / "skill.md"
HISTORY_PATH = ROOT / "history.json"
OUTPUT_ROOT = ROOT / "outputs"

KST = timezone(timedelta(hours=9))
FONT_KR = "맑은 고딕"
FONT_EN = "Arial"

MODEL = "claude-opus-4-7"
MAX_TOKENS = 16000

JOBS = [
    "영업", "마케팅", "재무회계", "인사", "생산관리",
    "IT개발", "MD바잉", "제약영업", "품질관리", "물류SCM", "홍보",
]
YEARS = ["6개월~1년차", "1~3년차", "4~6년차"]
SITUATIONS = [
    "같은직무이동", "같은직무+업계이동", "직무변동",
    "사기업→공기업", "공기업→사기업",
]

ROUTINE_JOBS = {"영업", "재무회계", "인사", "생산관리", "MD바잉", "품질관리", "물류SCM"}
PROJECT_JOBS = {"마케팅", "IT개발", "제약영업", "홍보"}

FORMAT_NAMES = {
    "F1": "업무단위·간단형",
    "F2": "업무단위·상세형",
    "F3": "프로젝트 단위형",
    "F4": "성과 강조 역순형",
    "F5": "AR/TAR 한 줄 요약형",
    "F6": "인턴/저연차 정성형",
}


# ===== history & selection =====

def init_history() -> dict:
    return {
        "current_cycle": 1,
        "remaining_in_cycle": list(JOBS),
        "used_combos": [],
        "history": [],
    }


def load_history() -> dict:
    if not HISTORY_PATH.exists():
        return init_history()
    return json.loads(HISTORY_PATH.read_text(encoding="utf-8"))


def save_history(h: dict) -> None:
    HISTORY_PATH.write_text(
        json.dumps(h, ensure_ascii=False, indent=2), encoding="utf-8"
    )


def pick_jobs(history: dict) -> list[str]:
    if len(history["remaining_in_cycle"]) < 3:
        history["current_cycle"] += 1
        history["remaining_in_cycle"] = list(JOBS)
    selected = random.sample(history["remaining_in_cycle"], 3)
    for j in selected:
        history["remaining_in_cycle"].remove(j)
    return selected


def pick_combo(history: dict, job: str) -> tuple[str, str]:
    used = set(history["used_combos"])
    candidates = [
        (y, s) for y in YEARS for s in SITUATIONS
        if f"{job}|{y}|{s}" not in used
    ]
    if not candidates:
        history["used_combos"] = [
            c for c in history["used_combos"] if not c.startswith(f"{job}|")
        ]
        candidates = [(y, s) for y in YEARS for s in SITUATIONS]
    year, situation = random.choice(candidates)
    history["used_combos"].append(f"{job}|{year}|{situation}")
    return year, situation


def get_format_candidates(job: str, year: str) -> list[str]:
    if year == "6개월~1년차":
        return ["F1", "F5", "F6"]
    cands = ["F1", "F4", "F5"]
    if job in ROUTINE_JOBS:
        cands.append("F2")
    if job in PROJECT_JOBS:
        cands.append("F3")
    return cands


def pick_formats(job: str, year: str, recent_history: list[dict]) -> list[str]:
    cands = get_format_candidates(job, year)
    if len(cands) < 2:
        return cands * 2 if cands else ["F1", "F5"]

    recent_pairs: set[tuple[str, str]] = set()
    for h in recent_history[-5:]:
        for o in h.get("outputs", []):
            if o["job"] == job and o["year"] == year:
                fmts = o.get("formats", [])
                if len(fmts) == 2:
                    recent_pairs.add(tuple(sorted(fmts)))

    for _ in range(20):
        pair = tuple(sorted(random.sample(cands, 2)))
        if pair not in recent_pairs:
            return list(pair)
    return list(random.sample(cands, 2))


# ===== Claude API =====

USER_TEMPLATE = """[오늘 작성할 직무 정보]
직무: {job}
연차: {year}
이직 상황: {situation}
양식 1: {f1} ({f1_name})
양식 2: {f2} ({f2_name})

[지시]
이 파일 PART B 의 모든 규칙(8원칙·양식 가이드·직무별 성과 표현·전략 패턴·docx 생성 규칙·체크리스트)을 적용하여, 위 직무에 대한 가상 인물 1명을 만들고, 해당 인물의 동일한 경력 정보를 양식 1·양식 2로 각각 작성하시오.

[가상 인물 생성 규칙]
- 이름: 한글 이름 (매번 다르게)
- 회사명: 실제 기업명 사용 금지 (○○패션, △△제약 등 마스킹 표기)
- 학력: OO대학교
- 근무기간: 연차에 맞게 오늘 날짜({today}) 기준으로 역산
- 양식 1과 양식 2는 동일 인물의 동일 경력을 다른 양식으로 표현 (직무·연차·회사명·근무기간·핵심 업무·핵심 성과 수치는 동일 유지)
- 이직 상황 반영: "{situation}"

[출력 형식 - 반드시 지킬 것]
첫 줄에 PERSONA_NAME 적고, 두 양식 본문을 각각 START/END 마커로 감싸 출력. 다른 메타텍스트·설명·code fence 절대 금지.

PERSONA_NAME: <한글 이름>
---FORMAT_{f1}_START---
<{f1} 양식 경력기술서 본문 마크다운 (본문만)>
---FORMAT_{f1}_END---
---FORMAT_{f2}_START---
<{f2} 양식 경력기술서 본문 마크다운 (본문만)>
---FORMAT_{f2}_END---

[본문 마크다운 규칙]
- 헤더 정보(회사명, 부서, 직급, 재직기간 등)는 줄바꿈된 텍스트 라인으로 작성
- 섹션 번호는 "1. 주요 업무" 형태
- 하위 카테고리는 "  가. 카테고리명", "  나. ..." 형태 (들여쓰기 2칸)
- 카테고리 하위 항목은 "    - 내용" 형태 (들여쓰기 4칸)
- 성과 요약은 "• 성과 요약 (볼드)" — 마크다운 볼드 **텍스트** 사용
- 성과 HOW 항목은 "    - 행동/전략" (대시) 또는 "    · 행동/전략" (가운데점)
- → 화살표 사용 금지
- 강조는 **텍스트** 만 사용 (이탤릭·헤더 # 사용 금지)
"""


def call_claude(skill_md: str, job: str, year: str, situation: str,
                f1: str, f2: str) -> str:
    client = anthropic.Anthropic()
    today = datetime.now(KST).strftime("%Y-%m-%d")
    user_msg = USER_TEMPLATE.format(
        job=job, year=year, situation=situation,
        f1=f1, f1_name=FORMAT_NAMES[f1],
        f2=f2, f2_name=FORMAT_NAMES[f2],
        today=today,
    )
    chunks: list[str] = []
    with client.messages.stream(
        model=MODEL,
        max_tokens=MAX_TOKENS,
        thinking={"type": "adaptive"},
        output_config={"effort": "high"},
        system=[
            {
                "type": "text",
                "text": skill_md,
                "cache_control": {"type": "ephemeral"},
            }
        ],
        messages=[{"role": "user", "content": user_msg}],
    ) as stream:
        for text in stream.text_stream:
            chunks.append(text)
        final = stream.get_final_message()
    print(
        f"[{job}] usage: in={final.usage.input_tokens}, "
        f"out={final.usage.output_tokens}, "
        f"cache_read={final.usage.cache_read_input_tokens}, "
        f"cache_create={final.usage.cache_creation_input_tokens}",
        file=sys.stderr,
    )
    return "".join(chunks)


def parse_response(text: str, f1: str, f2: str) -> tuple[str, str, str]:
    name_m = re.search(r"PERSONA_NAME\s*:\s*(.+)", text)
    name = name_m.group(1).strip() if name_m else "OOO"

    def extract(code: str) -> str:
        pattern = (
            rf"---FORMAT_{re.escape(code)}_START---\s*\n(.*?)\n"
            rf"---FORMAT_{re.escape(code)}_END---"
        )
        m = re.search(pattern, text, re.DOTALL)
        return m.group(1).strip() if m else ""

    return name, extract(f1), extract(f2)


# ===== docx rendering =====

def _set_run_font(run, size_pt, bold=False):
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), FONT_EN)
    rfonts.set(qn("w:hAnsi"), FONT_EN)
    rfonts.set(qn("w:eastAsia"), FONT_KR)
    rfonts.set(qn("w:cs"), FONT_EN)
    run.font.size = Pt(size_pt)
    run.bold = bold


def _set_paragraph_spacing(p, line=1.4, before=0, after=4):
    pf = p.paragraph_format
    pf.line_spacing = line
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)


def _add_runs_with_bold(p, text, size, base_bold=False):
    for part in re.split(r"(\*\*.+?\*\*)", text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            _set_run_font(run, size, bold=True)
        else:
            run = p.add_run(part)
            _set_run_font(run, size, bold=base_bold)


def _add_para(doc, text, size=11, bold=False, indent_cm=0.0,
              align=None, before=0, after=4):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if indent_cm > 0:
        p.paragraph_format.left_indent = Cm(indent_cm)
    _set_paragraph_spacing(p, line=1.4, before=before, after=after)
    _add_runs_with_bold(p, text, size, base_bold=bold)


def _add_blank(doc):
    p = doc.add_paragraph()
    _set_paragraph_spacing(p, line=1.0, after=0)


def _setup_document(doc):
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    normal = doc.styles["Normal"]
    normal.font.name = FONT_EN
    normal.font.size = Pt(11)
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), FONT_EN)
    rfonts.set(qn("w:hAnsi"), FONT_EN)
    rfonts.set(qn("w:eastAsia"), FONT_KR)


def render_to_docx(title: str, body_md: str, output_path: Path) -> None:
    doc = Document()
    _setup_document(doc)

    _add_para(doc, title, size=14, bold=True,
              align=WD_ALIGN_PARAGRAPH.CENTER, before=12, after=12)
    _add_blank(doc)

    for raw in body_md.split("\n"):
        line = raw.rstrip()
        stripped = line.strip()
        if not stripped:
            _add_blank(doc)
            continue

        leading = len(line) - len(line.lstrip(" "))
        level = leading // 2

        # Numbered section: "1. xxx" / "2. xxx"
        if re.match(r"^\d+\.\s+\S", stripped):
            _add_para(doc, stripped, size=11.5, bold=True, before=6, after=4)
        # Korean letter section: "가. xxx"
        elif re.match(r"^[가-힣]\.\s+\S", stripped):
            _add_para(doc, stripped, size=11, bold=True,
                      indent_cm=0.4 + 0.3 * max(0, level - 1))
        # Bracketed header: [...]
        elif stripped.startswith("[") and stripped.endswith("]"):
            _add_para(doc, stripped, size=11, bold=True, before=4, after=4)
        # 성과 요약 bullet: "• xxx"
        elif stripped.startswith("•"):
            indent = 0.4 + 0.3 * max(0, level - 1)
            _add_para(doc, stripped, size=11, bold=False, indent_cm=indent)
        # HOW sub bullets: "- xxx" or "· xxx"
        elif stripped.startswith("- ") or stripped.startswith("· "):
            indent = 0.8 + 0.3 * max(0, level - 1)
            _add_para(doc, stripped, size=10.5, indent_cm=indent)
        # Header info lines / plain
        else:
            _add_para(doc, stripped, size=11)

    doc.save(str(output_path))


# ===== entry =====

def main():
    if not os.environ.get("ANTHROPIC_API_KEY"):
        print("ERROR: ANTHROPIC_API_KEY env var not set", file=sys.stderr)
        sys.exit(1)

    if not SKILL_PATH.exists():
        print(f"ERROR: {SKILL_PATH} not found", file=sys.stderr)
        sys.exit(1)

    today = datetime.now(KST).strftime("%Y-%m-%d")
    today_dir = OUTPUT_ROOT / today

    if today_dir.exists() and any(today_dir.iterdir()):
        print(f"skip: {today_dir} already exists with files", file=sys.stderr)
        return

    history = load_history()
    if any(h.get("date") == today for h in history.get("history", [])):
        print(f"skip: history already has entry for {today}", file=sys.stderr)
        return

    skill_md = SKILL_PATH.read_text(encoding="utf-8")

    selected_jobs = pick_jobs(history)
    print(f"selected jobs: {selected_jobs}", file=sys.stderr)

    today_dir.mkdir(parents=True, exist_ok=True)
    today_outputs = []

    for job in selected_jobs:
        year, situation = pick_combo(history, job)
        formats = pick_formats(job, year, history.get("history", []))
        f1, f2 = formats[0], formats[1]
        print(
            f"[{job}] year={year}, situation={situation}, "
            f"formats=[{f1}, {f2}]",
            file=sys.stderr,
        )

        response = call_claude(skill_md, job, year, situation, f1, f2)
        name, body1, body2 = parse_response(response, f1, f2)
        if not body1 or not body2:
            print(
                f"WARN: missing body for {job} (body1={len(body1)}, "
                f"body2={len(body2)})",
                file=sys.stderr,
            )

        files: list[str] = []
        for fmt, body in [(f1, body1), (f2, body2)]:
            if not body:
                continue
            title = (
                f"경력기술서 - {job} ({year}, {fmt} {FORMAT_NAMES[fmt]})\n"
                f"{name} · {situation}"
            )
            filename = f"경력기술서_{job}_{fmt}.docx"
            output_path = today_dir / filename
            render_to_docx(title, body, output_path)
            files.append(str(output_path.relative_to(ROOT).as_posix()))
            print(f"saved: {output_path}", file=sys.stderr)

        today_outputs.append({
            "job": job,
            "year": year,
            "situation": situation,
            "formats": [f1, f2],
            "person_name": name,
            "files": files,
        })

    history.setdefault("history", []).append({
        "date": today,
        "cycle": history["current_cycle"],
        "outputs": today_outputs,
    })
    save_history(history)
    print(f"history.json updated for {today}", file=sys.stderr)


if __name__ == "__main__":
    main()
